"""
stripe_billing.py — Módulo de Billing Stripe
=============================================
Integração completa Stripe Subscriptions para FX Gestão de Frotas.

Uso no estudo_de_rede.py:
    from stripe_billing import (
        criar_checkout_session,
        processar_webhook,
        requer_plano,
        get_plano_atual,
        portal_cliente
    )
"""

import os
import stripe
import streamlit as st
from supabase import create_client

# ─────────────────────────────────────────────
# Configuração
# ─────────────────────────────────────────────
stripe.api_key = os.environ.get("STRIPE_SECRET_KEY")

SUPABASE_URL = os.environ.get("SUPABASE_URL", "")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "")
try:
    supabase = create_client(SUPABASE_URL, SUPABASE_KEY) if SUPABASE_URL and SUPABASE_KEY else None
except Exception:
    supabase = None

APP_URL = os.environ.get("APP_URL", "https://fxgestaodefrotasonline.com")

# Price IDs dos planos (Stripe)
PLANOS = {
    "gratuito":    {"price_id": None,                              "max_usuarios": 1,  "max_veiculos": 10},
    "basico":      {"price_id": "price_1Tf2ftRoRAomG8bP8hqJDGNC", "max_usuarios": 5,  "max_veiculos": 50},
    "profissional":{"price_id": "price_1Tf2gJRoRAomG8bPxUCKnKZy", "max_usuarios": 20, "max_veiculos": 200},
    "enterprise":  {"price_id": "price_1Tf2goRoRAomG8bPU40Yll9M", "max_usuarios": -1, "max_veiculos": -1},
}

# Mapeamento price_id → nome do plano
PRICE_ID_PARA_PLANO = {v["price_id"]: k for k, v in PLANOS.items() if v["price_id"]}

# Hierarquia dos planos (para comparação)
HIERARQUIA = {"gratuito": 0, "basico": 1, "profissional": 2, "enterprise": 3}


# ─────────────────────────────────────────────
# Checkout — redireciona para o Stripe
# ─────────────────────────────────────────────
def criar_checkout_session(plano: str, empresa_id: str, email: str) -> str:
    """
    Cria uma sessão de checkout no Stripe e retorna a URL.
    
    Uso:
        url = criar_checkout_session("basico", empresa_id, user_email)
        st.markdown(f'<meta http-equiv="refresh" content="0;url={url}">', unsafe_allow_html=True)
    """
    price_id = PLANOS[plano]["price_id"]
    if not price_id:
        raise ValueError(f"Plano '{plano}' não tem price_id configurado.")

    # Busca ou cria customer no Stripe
    empresa = supabase.table("empresas").select("stripe_customer_id").eq("id", empresa_id).single().execute()
    customer_id = empresa.data.get("stripe_customer_id")

    if not customer_id:
        customer = stripe.Customer.create(email=email, metadata={"empresa_id": empresa_id})
        customer_id = customer.id
        supabase.table("empresas").update({"stripe_customer_id": customer_id}).eq("id", empresa_id).execute()

    # Cria sessão de checkout
    session = stripe.checkout.Session.create(
        customer=customer_id,
        payment_method_types=["card"],
        line_items=[{"price": price_id, "quantity": 1}],
        mode="subscription",
        success_url=f"{APP_URL}?checkout=success&session_id={{CHECKOUT_SESSION_ID}}",
        cancel_url=f"{APP_URL}?checkout=cancelled",
        metadata={"empresa_id": empresa_id, "plano": plano},
        subscription_data={"metadata": {"empresa_id": empresa_id, "plano": plano}},
    )
    return session.url


# ─────────────────────────────────────────────
# Portal do Cliente — autoatendimento de billing
# ─────────────────────────────────────────────
def portal_cliente(empresa_id: str) -> str:
    """
    Retorna URL do portal Stripe onde o cliente pode:
    - Trocar cartão
    - Fazer upgrade/downgrade
    - Cancelar assinatura
    """
    empresa = supabase.table("empresas").select("stripe_customer_id").eq("id", empresa_id).single().execute()
    customer_id = empresa.data.get("stripe_customer_id")

    if not customer_id:
        raise ValueError("Empresa não tem customer_id no Stripe.")

    session = stripe.billing_portal.Session.create(
        customer=customer_id,
        return_url=APP_URL,
    )
    return session.url


# ─────────────────────────────────────────────
# Webhooks — processa eventos do Stripe
# ─────────────────────────────────────────────
def processar_webhook(payload: bytes, sig_header: str) -> dict:
    """
    Valida e processa um evento webhook do Stripe.
    Garante idempotência verificando stripe_events.
    
    Retorna: {"status": "ok"|"ignorado"|"erro", "mensagem": str}
    """
    webhook_secret = os.environ.get("STRIPE_WEBHOOK_SECRET")

    try:
        evento = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
    except stripe.error.SignatureVerificationError:
        return {"status": "erro", "mensagem": "Assinatura inválida"}

    event_id   = evento["id"]
    event_tipo = evento["type"]

    # ── Idempotência: ignora evento já processado ──
    existente = supabase.table("stripe_events").select("id").eq("stripe_event_id", event_id).execute()
    if existente.data:
        return {"status": "ignorado", "mensagem": f"Evento {event_id} já processado"}

    # ── Salva o evento ──
    supabase.table("stripe_events").insert({
        "stripe_event_id": event_id,
        "tipo": event_tipo,
        "payload": evento,
    }).execute()

    # ── Processa por tipo ──
    dados = evento["data"]["object"]

    if event_tipo == "checkout.session.completed":
        _handle_checkout_completed(dados)

    elif event_tipo == "invoice.payment_succeeded":
        _handle_payment_succeeded(dados)

    elif event_tipo == "invoice.payment_failed":
        _handle_payment_failed(dados)

    elif event_tipo == "customer.subscription.deleted":
        _handle_subscription_deleted(dados)

    elif event_tipo == "customer.subscription.updated":
        _handle_subscription_updated(dados)

    # ── Marca como processado ──
    supabase.table("stripe_events").update({"processado_em": "now()"}).eq("stripe_event_id", event_id).execute()

    return {"status": "ok", "mensagem": f"Evento {event_tipo} processado"}


def _handle_checkout_completed(session: dict):
    try:
        _webhook_pos_pagamento(session)
    except Exception:
        pass
    """Ativa tenant após checkout bem-sucedido."""
    empresa_id       = session.get("metadata", {}).get("empresa_id")
    plano            = session.get("metadata", {}).get("plano", "basico")
    subscription_id  = session.get("subscription")
    customer_id      = session.get("customer")

    if not empresa_id:
        return

    supabase.table("empresas").update({
        "status":                  "ativo",
        "plano":                   plano,
        "stripe_subscription_id":  subscription_id,
        "stripe_customer_id":      customer_id,
        "max_usuarios":            PLANOS[plano]["max_usuarios"],
        "max_veiculos":            PLANOS[plano]["max_veiculos"],
    }).eq("id", empresa_id).execute()


def _handle_payment_succeeded(invoice: dict):
    """Renova status ativo e registra fatura."""
    customer_id = invoice.get("customer")
    empresa     = supabase.table("empresas").select("id").eq("stripe_customer_id", customer_id).execute()
    if not empresa.data:
        return

    empresa_id = empresa.data[0]["id"]

    supabase.table("empresas").update({"status": "ativo"}).eq("id", empresa_id).execute()

    # Registra fatura
    supabase.table("invoices").insert({
        "empresa_id":       empresa_id,
        "stripe_invoice_id": invoice.get("id"),
        "valor_cents":      invoice.get("amount_paid", 0),
        "status":           "pago",
        "periodo_inicio":   _ts(invoice.get("period_start")),
        "periodo_fim":      _ts(invoice.get("period_end")),
    }).execute()


def _handle_payment_failed(invoice: dict):
    """Suspende tenant após falha de pagamento."""
    customer_id = invoice.get("customer")
    empresa     = supabase.table("empresas").select("id").eq("stripe_customer_id", customer_id).execute()
    if not empresa.data:
        return

    empresa_id = empresa.data[0]["id"]
    supabase.table("empresas").update({"status": "suspenso"}).eq("id", empresa_id).execute()

    supabase.table("invoices").insert({
        "empresa_id":        empresa_id,
        "stripe_invoice_id": invoice.get("id"),
        "valor_cents":       invoice.get("amount_due", 0),
        "status":            "falhou",
        "periodo_inicio":    _ts(invoice.get("period_start")),
        "periodo_fim":       _ts(invoice.get("period_end")),
    }).execute()


def _handle_subscription_deleted(subscription: dict):
    """Cancela tenant quando assinatura é deletada."""
    customer_id = subscription.get("customer")
    empresa     = supabase.table("empresas").select("id").eq("stripe_customer_id", customer_id).execute()
    if not empresa.data:
        return

    supabase.table("empresas").update({
        "status":       "cancelado",
        "cancelado_em": "now()",
        "plano":        "gratuito",
    }).eq("id", empresa.data[0]["id"]).execute()


def _handle_subscription_updated(subscription: dict):
    """Atualiza plano após upgrade/downgrade."""
    customer_id = subscription.get("customer")
    empresa     = supabase.table("empresas").select("id").eq("stripe_customer_id", customer_id).execute()
    if not empresa.data:
        return

    empresa_id = empresa.data[0]["id"]

    # Descobre qual plano pelo price_id
    items    = subscription.get("items", {}).get("data", [])
    price_id = items[0]["price"]["id"] if items else None
    plano    = PRICE_ID_PARA_PLANO.get(price_id, "basico")

    supabase.table("empresas").update({
        "plano":        plano,
        "status":       subscription.get("status", "ativo"),
        "max_usuarios": PLANOS[plano]["max_usuarios"],
        "max_veiculos": PLANOS[plano]["max_veiculos"],
    }).eq("id", empresa_id).execute()


# ─────────────────────────────────────────────
# Feature Flags — controle de acesso por plano
# ─────────────────────────────────────────────
def get_plano_atual() -> str:
    """Retorna o plano atual do tenant logado."""
    return st.session_state.get("tenant_plano", "gratuito")


def requer_plano(plano_minimo: str):
    """
    Bloqueia a execução se o tenant não tiver o plano mínimo.
    
    Uso:
        requer_plano("profissional")
        # código abaixo só executa para Pro e Enterprise
    """
    plano_atual = get_plano_atual()
    empresa_id  = st.session_state.get("empresa_id")

    if HIERARQUIA.get(plano_atual, 0) < HIERARQUIA.get(plano_minimo, 0):
        nomes = {"basico": "Básico (R$149/mês)", "profissional": "Pro (R$349/mês)", "enterprise": "Enterprise (R$899/mês)"}
        st.warning(f"🔒 Esta função requer o plano **{nomes.get(plano_minimo, plano_minimo)}**.")

        col1, col2 = st.columns(2)
        with col1:
            if st.button("🚀 Fazer upgrade agora", type="primary"):
                if empresa_id:
                    email = st.session_state.get("user_email", "")
                    url   = criar_checkout_session(plano_minimo, empresa_id, email)
                    st.markdown(f'<meta http-equiv="refresh" content="0;url={url}">', unsafe_allow_html=True)
        with col2:
            st.caption(f"Plano atual: **{plano_atual.capitalize()}**")

        st.stop()


def verificar_limite(recurso: str, quantidade_atual: int) -> bool:
    """
    Verifica se o tenant atingiu o limite do plano.
    
    Uso:
        if not verificar_limite("veiculos", total_veiculos):
            st.error("Limite de veículos atingido. Faça upgrade!")
            st.stop()
    """
    plano_atual = get_plano_atual()
    limites     = PLANOS.get(plano_atual, PLANOS["gratuito"])

    if recurso == "veiculos":
        limite = limites["max_veiculos"]
    elif recurso == "usuarios":
        limite = limites["max_usuarios"]
    else:
        return True

    if limite == -1:  # ilimitado (Enterprise)
        return True

    return quantidade_atual < limite


# ─────────────────────────────────────────────
# UI — Tela de Planos
# ─────────────────────────────────────────────
def mostrar_tela_planos():
    """Renderiza a tela de seleção de planos com integração real ao Stripe."""
    import streamlit as _st

    # ── Resolve empresa e email do usuário logado ──────────────────
    _emp_pl    = _st.session_state.get("_empresa_ativa") or {}
    empresa_id = _emp_pl.get("id", "")
    plano_atual= _emp_pl.get("plano", "gratuito")
    _auth_user = _st.session_state.get("_auth_user") or {}
    email      = (_auth_user.get("email") or
                  _st.session_state.get("_auth_email", "") or "")

    _st.markdown(
        "<h2 style='margin:0 0 4px'>🚀 Planos & Assinatura</h2>"
        "<p style='color:#666;margin:0 0 20px'>Escolha o plano ideal para sua frota.</p>",
        unsafe_allow_html=True,
    )

    if not empresa_id:
        _st.warning("Empresa não identificada. Faça login novamente.")
        return

    PLANOS_INFO = [
        {
            "key": "gratuito",
            "icon": "🆓", "nome": "Gratuito", "preco": "R$ 0/mês",
            "recursos": ["1 usuário", "10 veículos", "Consulta de postos ANP"],
        },
        {
            "key": "basico",
            "icon": "📦", "nome": "Básico", "preco": "R$ 149/mês",
            "recursos": ["5 usuários", "50 veículos", "Exportação Excel", "Suporte por e-mail"],
        },
        {
            "key": "profissional",
            "icon": "⭐", "nome": "Pro", "preco": "R$ 349/mês",
            "recursos": ["20 usuários", "200 veículos", "Relatórios avançados", "API REST", "Assistente IA"],
        },
        {
            "key": "enterprise",
            "icon": "🏢", "nome": "Enterprise", "preco": "R$ 899/mês",
            "recursos": ["Ilimitado", "SSO / SAML", "SLA 99,95%", "Suporte dedicado", "Onboarding guiado"],
        },
    ]

    cols = _st.columns(4)
    for col, plano in zip(cols, PLANOS_INFO):
        with col:
            is_atual = (plano["key"] == plano_atual)
            borda    = "2px solid #1565c0" if is_atual else "1px solid #e0e0e0"
            _st.markdown(
                f"<div style='border:{borda};border-radius:12px;padding:16px 12px;"
                f"background:{'#f0f7ff' if is_atual else '#fff'};min-height:220px'>"
                f"<div style='font-size:24px'>{plano['icon']}</div>"
                f"<div style='font-weight:700;font-size:15px;margin:4px 0'>{plano['nome']}</div>"
                f"<div style='color:#1565c0;font-weight:600;margin-bottom:10px'>{plano['preco']}</div>"
                + "".join(f"<div style='font-size:12px;color:#555;margin-bottom:2px'>✔ {r}</div>"
                          for r in plano["recursos"])
                + "</div>",
                unsafe_allow_html=True,
            )
            _st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

            if is_atual:
                _st.button("✅ Plano atual", disabled=True,
                           key=f"btn_plano_{plano['key']}", use_container_width=True)
            elif plano["key"] == "gratuito":
                _st.button("Fazer downgrade", disabled=True,
                           key=f"btn_plano_{plano['key']}", use_container_width=True)
            else:
                if _st.button(f"Assinar {plano['nome']}",
                              key=f"btn_plano_{plano['key']}",
                              type="primary", use_container_width=True):
                    try:
                        with _st.spinner("Preparando checkout seguro..."):
                            _url = criar_checkout_session(plano["key"], empresa_id, email)
                        _st.markdown(
                            f"<div style='background:#e8f5e9;border:1px solid #a5d6a7;"
                            f"border-radius:8px;padding:12px 16px;margin-top:8px'>"
                            f"✅ <b>Checkout criado!</b><br>"
                            f"<a href='{_url}' target='_blank' style='color:#1565c0;font-weight:600'>"
                            f"👉 Clique aqui para pagar com segurança no Stripe</a></div>",
                            unsafe_allow_html=True,
                        )
                        _st.link_button(
                            f"💳 Ir para pagamento — {plano['nome']}",
                            url=_url, use_container_width=True,
                        )
                    except Exception as _e_str:
                        _st.error(f"Erro ao criar sessão de pagamento: {_e_str}")

    _st.divider()
    _st.markdown("#### ⚙️ Gerenciar assinatura existente")
    _st.caption("Altere forma de pagamento, cancele ou veja faturas anteriores.")
    if _st.button("Acessar portal do cliente Stripe", key="btn_portal_stripe"):
        try:
            with _st.spinner("Abrindo portal..."):
                _url_portal = portal_cliente(empresa_id)
            _st.link_button("🔗 Acessar portal", url=_url_portal)
        except Exception as _e_portal:
            _st.error(f"Erro ao abrir portal: {_e_portal}")

def _ts(unix_timestamp):
    """Converte timestamp Unix para ISO 8601."""
    if not unix_timestamp:
        return None
    from datetime import datetime, timezone
    return datetime.fromtimestamp(unix_timestamp, tz=timezone.utc).isoformat()

# BLOCO_IMPORTS_MELHORIAS_V1
import smtplib, ssl, hashlib, io, tempfile
from email.mime.multipart import MIMEMultipart
from email.mime.text       import MIMEText
from email.mime.base       import MIMEBase
from email                 import encoders
import datetime as _dt_termo

_TERMO_GITHUB_URL = (
    "https://raw.githubusercontent.com/dperuffo/estudo-de-rede/"
    "master/estudo-de-rede/Termo_Adesao_FNI_Gestao_Frotas.docx"
)
_SMTP_HOST  = "smtp.hostinger.com"
_SMTP_PORT  = 465
_SMTP_USER  = "contato@fxgestaodefrotasonline.com"
_SMTP_PASS  = os.environ.get("SMTP_PASSWORD", "")
_EMAIL_FROM = "FNI Gestão de Frotas <contato@fxgestaodefrotasonline.com>"

@st.cache_data(ttl=300)
def _baixar_termo_docx() -> bytes:
    try:
        req = urllib.request.Request(
            _TERMO_GITHUB_URL, headers={"User-Agent": "FNI-App/1.0"}
        )
        with urllib.request.urlopen(req, timeout=15) as r:
            return r.read()
    except Exception:
        return b""

def _docx_para_html(docx_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        linhas = []
        for para in doc.paragraphs:
            txt = para.text.strip()
            if not txt:
                linhas.append("<br>")
                continue
            style = para.style.name.lower() if para.style else ""
            if "heading 1" in style or "titulo 1" in style:
                linhas.append("<h3 style=\"color:#0D1B3E;margin:16px 0 6px\">" + txt + "</h3>")
            elif "heading" in style or "titulo" in style:
                linhas.append("<h4 style=\"color:#1565C0;margin:12px 0 4px\">" + txt + "</h4>")
            else:
                linhas.append("<p style=\"margin:4px 0;font-size:13px;line-height:1.6\">" + txt + "</p>")
        return "\n".join(linhas)
    except Exception:
        return "<p>Nao foi possivel renderizar o termo. Faca o download para visualizar.</p>"

def _gerar_hash_termo(docx_bytes: bytes) -> str:
    return hashlib.sha256(docx_bytes).hexdigest()

def _registrar_aceite_termo(email: str, plano: str, hash_termo: str,
                             empresa_id: str = "") -> str | None:
    try:
        from supabase import create_client
        _url = os.environ.get("SUPABASE_URL", "")
        _key = os.environ.get("SUPABASE_KEY", "")
        if not (_url and _key):
            return None
        db  = create_client(_url, _key)
        ts  = _dt_termo.datetime.now(
            _dt_termo.timezone(_dt_termo.timedelta(hours=-3))
        ).isoformat()
        resp = db.table("termos_aceite").insert({
            "email":       email.lower().strip(),
            "plano":       plano,
            "hash_termo":  hash_termo,
            "aceito_em":   ts,
            "empresa_id":  empresa_id or None,
            "versao_termo": "1.0",
        }).execute()
        return (resp.data or [{}])[0].get("id")
    except Exception:
        return None

def _gerar_pdf_termo_assinado(docx_bytes: bytes, email: str, plano: str,
                               aceito_em: str, hash_termo: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles    import ParagraphStyle
        from reportlab.lib.units     import cm
        from reportlab.lib           import colors
        from reportlab.platypus      import (SimpleDocTemplate, Paragraph,
                                              Spacer, HRFlowable,
                                              Table, TableStyle)
        buf     = io.BytesIO()
        doc_pdf = SimpleDocTemplate(buf, pagesize=A4,
                                    rightMargin=2*cm, leftMargin=2*cm,
                                    topMargin=2*cm, bottomMargin=2*cm)
        story   = []

        s_titulo = ParagraphStyle("titulo", fontSize=18,
                                  fontName="Helvetica-Bold",
                                  textColor=colors.HexColor("#0D1B3E"),
                                  spaceAfter=4)
        s_sub    = ParagraphStyle("sub", fontSize=10, fontName="Helvetica",
                                  textColor=colors.HexColor("#1565C0"),
                                  spaceAfter=12)
        s_body   = ParagraphStyle("body", fontSize=9, fontName="Helvetica",
                                  leading=14, spaceAfter=6)

        story.append(Paragraph("FNI Gestao de Frotas", s_titulo))
        story.append(Paragraph("Termo de Adesao - Copia do Assinante", s_sub))
        story.append(HRFlowable(width="100%", thickness=2,
                                color=colors.HexColor("#0D1B3E")))
        story.append(Spacer(1, 0.4*cm))

        try:
            from docx import Document
            tdoc = Document(io.BytesIO(docx_bytes))
            for para in tdoc.paragraphs:
                txt = para.text.strip()
                if not txt:
                    story.append(Spacer(1, 0.2*cm))
                    continue
                sn = para.style.name.lower() if para.style else ""
                if "heading" in sn or "titulo" in sn:
                    sh = ParagraphStyle("h", fontSize=11,
                                        fontName="Helvetica-Bold",
                                        textColor=colors.HexColor("#0D1B3E"),
                                        spaceBefore=8, spaceAfter=4)
                else:
                    sh = s_body
                story.append(Paragraph(
                    txt.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"), sh
                ))
        except Exception:
            story.append(Paragraph("Termo conforme documento original.", s_body))

        story.append(Spacer(1, 0.6*cm))
        story.append(HRFlowable(width="100%", thickness=1,
                                color=colors.HexColor("#4FC3F7")))
        story.append(Spacer(1, 0.4*cm))

        s_ah  = ParagraphStyle("ah", fontSize=11, fontName="Helvetica-Bold",
                               textColor=colors.HexColor("#1B5E20"), spaceAfter=8)
        story.append(Paragraph("ASSINATURA ELETRONICA", s_ah))

        dados = [
            ["Assinante (e-mail):", email],
            ["Plano contratado:",   plano.upper()],
            ["Data e hora (BRT):",  aceito_em],
            ["Hash do documento:",  hash_termo[:32] + "..."],
            ["Validade juridica:",  "MP 2.200-2/2001 - ICP-Brasil"],
        ]
        tbl = Table(dados, colWidths=[4.5*cm, 12.5*cm])
        tbl.setStyle(TableStyle([
            ("FONTNAME",  (0,0), (-1,-1), "Helvetica"),
            ("FONTSIZE",  (0,0), (-1,-1), 8),
            ("FONTNAME",  (0,0), (0,-1),  "Helvetica-Bold"),
            ("TEXTCOLOR", (0,0), (0,-1),  colors.HexColor("#0D1B3E")),
            ("ROWBACKGROUNDS", (0,0), (-1,-1),
             [colors.HexColor("#F5F5F5"), colors.white]),
            ("GRID",     (0,0), (-1,-1), 0.3, colors.HexColor("#CCCCCC")),
            ("PADDING",  (0,0), (-1,-1), 5),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 0.4*cm))
        s_rod = ParagraphStyle("rod", fontSize=7, fontName="Helvetica",
                               textColor=colors.grey, leading=10)
        story.append(Paragraph(
            "Este documento e uma copia do Termo de Adesao com registro de aceite "
            "eletronico conforme a Medida Provisoria 2.200-2/2001. "
            "O hash SHA-256 garante sua integridade.", s_rod
        ))
        doc_pdf.build(story)
        return buf.getvalue()
    except Exception:
        return b""

def _enviar_email_termo(email_dest: str, plano: str,
                         pdf_bytes: bytes, aceito_em: str) -> bool:
    try:
        if not _SMTP_PASS:
            return False
        msg = MIMEMultipart()
        msg["From"]    = _EMAIL_FROM
        msg["To"]      = email_dest
        msg["Subject"] = "Seu Termo de Adesao FNI - Plano " + plano.capitalize()
        corpo = (
            "<html><body style=\"font-family:Arial,sans-serif;color:#222;"
            "max-width:600px;margin:0 auto\">"
            "<div style=\"background:#0D1B3E;padding:24px 32px;"
            "border-radius:8px 8px 0 0\">"
            "<h1 style=\"color:#fff;margin:0;font-size:22px\">FNI Gestao de Frotas</h1>"
            "<p style=\"color:#4FC3F7;margin:4px 0 0;font-size:13px\">"
            "Confirmacao de Adesao</p></div>"
            "<div style=\"background:#f9f9f9;padding:24px 32px;"
            "border:1px solid #e0e0e0\">"
            "<h2 style=\"color:#0D1B3E;font-size:17px\">"
            "Parabens pela sua assinatura!</h2>"
            "<p>Seu <b>Termo de Adesao ao Plano " + plano.capitalize() + "</b>"
            " foi registrado com sucesso.</p>"
            "<table style=\"width:100%;border-collapse:collapse;margin:16px 0\">"
            "<tr style=\"background:#EEF2FF\">"
            "<td style=\"padding:8px 12px;font-weight:bold;color:#0D1B3E;width:40%\">"
            "Plano contratado</td>"
            "<td style=\"padding:8px 12px\">" + plano.upper() + "</td></tr>"
            "<tr><td style=\"padding:8px 12px;font-weight:bold;color:#0D1B3E\">"
            "Data e hora (BRT)</td>"
            "<td style=\"padding:8px 12px\">" + aceito_em + "</td></tr>"
            "<tr style=\"background:#EEF2FF\">"
            "<td style=\"padding:8px 12px;font-weight:bold;color:#0D1B3E\">"
            "Assinatura eletronica</td>"
            "<td style=\"padding:8px 12px\">"
            "Registrada conforme MP 2.200-2/2001</td></tr></table>"
            "<p>Em anexo: <b>copia do Termo de Adesao</b> com o registro "
            "da sua assinatura eletronica.</p>"
            "<hr style=\"border:none;border-top:1px solid #ddd;margin:20px 0\">"
            "<p style=\"font-size:12px;color:#777\">"
            "FNI Gestao de Frotas - contato@fxgestaodefrotasonline.com</p>"
            "</div></body></html>"
        )
        msg.attach(MIMEText(corpo, "html", "utf-8"))
        if pdf_bytes:
            parte = MIMEBase("application", "octet-stream")
            parte.set_payload(pdf_bytes)
            encoders.encode_base64(parte)
            nome_pdf = "Termo_Adesao_FNI_" + plano.capitalize() + ".pdf"
            parte.add_header("Content-Disposition",
                             "attachment; filename=\"" + nome_pdf + "\"")
            msg.attach(parte)
        ctx = ssl.create_default_context()
        with smtplib.SMTP_SSL(_SMTP_HOST, _SMTP_PORT, context=ctx) as srv:
            srv.login(_SMTP_USER, _SMTP_PASS)
            srv.sendmail(_SMTP_USER, email_dest, msg.as_string())
        return True
    except Exception:
        return False

def _tela_termo_adesao(plano: str, preco: str,
                        email: str, empresa_id: str = "") -> bool:
    st.markdown(
        "<style>"
        ".fni-th{background:linear-gradient(135deg,#0D1B3E 0%,#1565C0 100%);"
        "border-radius:12px;padding:28px 32px;margin-bottom:24px;text-align:center}"
        ".fni-th h1{color:#fff;font-size:26px;margin:0 0 6px;font-weight:900}"
        ".fni-th p{color:#90CAF9;font-size:14px;margin:0}"
        ".fni-tb{background:#FAFAFA;border:1px solid #E0E0E0;border-radius:10px;"
        "padding:24px;max-height:420px;overflow-y:auto;margin-bottom:20px;"
        "font-size:13px;line-height:1.7;color:#333}"
        ".fni-pb{display:inline-block;background:#1565C0;color:#fff;"
        "border-radius:20px;padding:4px 16px;font-size:13px;"
        "font-weight:700;margin-bottom:16px}"
        "</style>",
        unsafe_allow_html=True
    )
    st.markdown(
        "<div class=\"fni-th\"><h1>Termo de Adesao</h1>"
        "<p>Leia atentamente antes de prosseguir para o pagamento</p></div>"
        "<div style=\"text-align:center;margin-bottom:16px\">"
        "<span class=\"fni-pb\">Plano " + plano.capitalize() +
        " - R$ " + preco + "/mes</span></div>",
        unsafe_allow_html=True
    )
    docx_bytes = _baixar_termo_docx()
    if not docx_bytes:
        st.error("Nao foi possivel carregar o termo. Tente novamente.")
        return False
    hash_termo = _gerar_hash_termo(docx_bytes)
    html_termo = _docx_para_html(docx_bytes)
    st.markdown(
        "<div class=\"fni-tb\">" + html_termo + "</div>",
        unsafe_allow_html=True
    )
    st.download_button(
        "Baixar Termo (.docx)",
        data=docx_bytes,
        file_name="Termo_Adesao_FNI_Gestao_Frotas.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    st.markdown("<div style=\"height:12px\"></div>", unsafe_allow_html=True)
    aceito = st.checkbox(
        "Li, compreendi e aceito integralmente os termos e condicoes acima.",
        key="chk_aceite_termo"
    )
    col_v, col_c = st.columns([1, 1])
    with col_v:
        if st.button("Voltar", use_container_width=True, key="btn_voltar_termo"):
            st.session_state.pop("_termo_plano", None)
            st.rerun()
    with col_c:
        if st.button("Continuar para Pagamento",
                     use_container_width=True,
                     disabled=not aceito,
                     type="primary",
                     key="btn_confirmar_termo"):
            aceito_em = _dt_termo.datetime.now(
                _dt_termo.timezone(_dt_termo.timedelta(hours=-3))
            ).strftime("%d/%m/%Y %H:%M:%S BRT")
            aceite_id = _registrar_aceite_termo(
                email=email, plano=plano,
                hash_termo=hash_termo, empresa_id=empresa_id
            )
            st.session_state["_termo_aceito"]    = True
            st.session_state["_termo_hash"]      = hash_termo
            st.session_state["_termo_aceito_em"] = aceito_em
            st.session_state["_termo_aceite_id"] = aceite_id
            st.session_state["_termo_docx"]      = docx_bytes
            st.rerun()
            return True
    return False

def _tela_transicao_fni(plano: str, preco: str, url_checkout: str):
    st.markdown(
        "<style>"
        ".fpw{max-width:520px;margin:0 auto;text-align:center;padding:32px 16px}"
        ".fpl{background:#0D1B3E;border-radius:16px;padding:28px 40px;"
        "display:inline-block;margin-bottom:28px;"
        "box-shadow:0 8px 32px rgba(13,27,62,0.35)}"
        ".fpl h1{color:#fff;font-size:32px;font-weight:900;margin:0 0 4px;"
        "letter-spacing:2px}"
        ".fpl p{color:#4FC3F7;font-size:13px;margin:0;letter-spacing:0.5px}"
        ".fpl hr{border:none;border-top:2px solid #4FC3F7;opacity:0.5;margin:10px 0 0}"
        ".fpc{background:#fff;border:1px solid #E3EAF6;border-radius:14px;"
        "padding:28px 32px;margin-top:8px;"
        "box-shadow:0 4px 20px rgba(13,27,62,0.08)}"
        ".fpn{font-size:22px;font-weight:900;color:#0D1B3E;margin:0 0 4px}"
        ".fpp{font-size:36px;font-weight:900;color:#1565C0;margin:0 0 16px}"
        ".fpb{display:inline-flex;align-items:center;gap:6px;"
        "background:#E8F5E9;color:#1B5E20;border-radius:20px;"
        "padding:4px 14px;font-size:12px;font-weight:700;margin-bottom:20px}"
        ".fps{background:#F0F4FF;border-radius:10px;padding:12px 16px;"
        "margin-bottom:20px;font-size:12px;color:#555;line-height:1.6}"
        "</style>",
        unsafe_allow_html=True
    )
    st.markdown(
        "<div class=\"fpw\">"
        "<div class=\"fpl\"><h1>FNI</h1><p>Gestao de Frotas</p><hr></div>"
        "<div class=\"fpc\">"
        "<p class=\"fpn\">Plano " + plano.capitalize() + "</p>"
        "<p class=\"fpp\">R$ " + preco +
        "<span style=\"font-size:14px;color:#888;font-weight:400\">/mes</span></p>"
        "<div class=\"fpb\">Termo de Adesao aceito</div>"
        "<div class=\"fps\"><b>Pagamento 100% seguro</b><br>"
        "Seus dados sao protegidos por criptografia SSL e processados pela "
        "<b>Stripe</b>, certificada PCI DSS nivel 1 - o mais alto padrao "
        "de seguranca para pagamentos online.</div>"
        "</div></div>",
        unsafe_allow_html=True
    )
    st.markdown(
        "<div style=\"max-width:520px;margin:16px auto 0\">",
        unsafe_allow_html=True
    )
    st.link_button(
        "Ir para Pagamento Seguro - Stripe",
        url=url_checkout,
        use_container_width=True,
        type="primary",
    )
    if st.button("Voltar", use_container_width=True,
                 key="btn_voltar_transicao"):
        st.session_state.pop("_termo_aceito", None)
        st.session_state.pop("_checkout_url", None)
        st.rerun()
    st.markdown(
        "<p style=\"text-align:center;font-size:11px;color:#aaa;margin-top:10px\">"
        "Powered by Stripe - Seus dados nunca sao armazenados pela FNI</p>"
        "</div>",
        unsafe_allow_html=True
    )

def _mostrar_tela_planos_com_termo():
    _email_usr  = (
        (st.session_state.get("_auth_user") or {}).get("email", "")
        or st.session_state.get("_auth_usuario_db", {}).get("email", "")
    )
    _empresa_id = st.session_state.get("_empresa_ativa", {}).get("id", "") or ""
    _plano_sel  = st.session_state.get("_termo_plano", "")
    _preco_sel  = st.session_state.get("_termo_preco", "")
    _aceito     = st.session_state.get("_termo_aceito", False)
    _url_chk    = st.session_state.get("_checkout_url", "")
    if _aceito and _url_chk:
        _tela_transicao_fni(_plano_sel, _preco_sel, _url_chk)
        return
    if _plano_sel and not _aceito:
        _tela_termo_adesao(
            plano=_plano_sel, preco=_preco_sel,
            email=_email_usr, empresa_id=_empresa_id
        )
        return
    _orig_checkout = None
    try:
        import stripe as _stripe
        _orig_checkout = _stripe.checkout.Session.create
        def _patched_checkout(**kwargs):
            sess  = _orig_checkout(**kwargs)
            meta  = kwargs.get("metadata", {})
            st.session_state["_checkout_url"] = sess.url
            st.session_state["_termo_plano"]  = meta.get("plano", "")
            st.session_state["_termo_preco"]  = meta.get("preco", "")
            return sess
        _stripe.checkout.Session.create = _patched_checkout
    except Exception:
        pass
    mostrar_tela_planos()
    try:
        if _orig_checkout:
            _stripe.checkout.Session.create = _orig_checkout
    except Exception:
        pass
    if (st.session_state.get("_checkout_url")
            and st.session_state.get("_termo_plano")
            and not st.session_state.get("_termo_aceito")):
        st.rerun()

def _webhook_pos_pagamento(session: dict):
    try:
        email = (
            session.get("customer_email")
            or (session.get("customer_details") or {}).get("email", "")
        )
        meta  = session.get("metadata") or {}
        plano = meta.get("plano", "assinatura")
        if not email:
            return
        from supabase import create_client
        _url = os.environ.get("SUPABASE_URL", "")
        _key = os.environ.get("SUPABASE_KEY", "")
        if not (_url and _key):
            return
        db   = create_client(_url, _key)
        rows = (
            db.table("termos_aceite")
            .select("hash_termo,aceito_em,id")
            .eq("email", email.lower())
            .eq("plano", plano)
            .order("aceito_em", desc=True)
            .limit(1)
            .execute()
        ).data or []
        hash_termo = rows[0]["hash_termo"] if rows else ""
        aceito_em  = rows[0]["aceito_em"]  if rows else ""
        docx_bytes = _baixar_termo_docx()
        pdf_bytes  = (
            _gerar_pdf_termo_assinado(
                docx_bytes, email, plano, aceito_em, hash_termo
            ) if docx_bytes else b""
        )
        ok = _enviar_email_termo(email, plano, pdf_bytes, aceito_em)
        if rows:
            db.table("termos_aceite").update({
                "pagamento_confirmado": True,
                "stripe_session_id":   session.get("id", ""),
                "email_enviado":       ok,
            }).eq("id", rows[0]["id"]).execute()
    except Exception:
        pass
